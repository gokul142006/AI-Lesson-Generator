"""
doc_export.py
----------------
Generates a professionally formatted Microsoft Word document (.docx)
for a lesson plan, using python-docx. Returns raw bytes so the
Streamlit layer can feed them straight into st.download_button.
"""

from __future__ import annotations

import io
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from utils import SECTION_SCHEMA, markdown_to_plain_lines

try:
    from excel_export import ExportError
except ImportError:  # pragma: no cover
    class ExportError(Exception):
        def __init__(self, user_message: str):
            super().__init__(user_message)
            self.user_message = user_message


PRIMARY_COLOR = RGBColor(0x1F, 0x3B, 0x57)


def _add_heading(doc: Document, text: str, size: int = 14):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = PRIMARY_COLOR
    heading.space_before = Pt(12)
    heading.space_after = Pt(4)
    return heading


def build_docx(lesson: Dict[str, Any], meta: Dict[str, Any]) -> bytes:
    """Build the Word document in memory and return its bytes."""
    try:
        doc = Document()

        # --- Title -----------------------------------------------------------
        title = lesson.get("lesson_title") or "Lesson Plan"
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = PRIMARY_COLOR

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("AI-Generated Lesson Plan")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(11)

        doc.add_paragraph()  # spacer

        # --- Metadata table ----------------------------------------------------
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light Grid Accent 1"
        meta_rows = [
            ("Subject", meta.get("subject", "")),
            ("Grade", meta.get("grade", "")),
            ("Topic", meta.get("topic", "")),
            ("Duration", meta.get("duration", "")),
        ]
        for i, (label, value) in enumerate(meta_rows):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()  # spacer

        # --- Sections -----------------------------------------------------------
        for key, label in SECTION_SCHEMA:
            if key == "lesson_title":
                continue
            content = lesson.get(key, "")
            if not content:
                continue

            _add_heading(doc, label)
            for line in markdown_to_plain_lines(content):
                if not line.strip():
                    continue
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.add_run(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as exc:
        raise ExportError(f"Could not generate the Word document: {exc}")
